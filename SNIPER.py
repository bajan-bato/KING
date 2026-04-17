#!/usr/bin/env python3
"""
Match EL/ID/otpremnica from first Excel, join with otpremnica/serijski from second Excel,
fill Word templates (second table, column "Serijski broj") with serial numbers.
Outputs joined.xlsx and flagged rows (with serial numbers added), both nicely formatted.
All spaces are removed from EL values for consistent matching.
"""

import os
import re
import pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from collections import defaultdict
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from openpyxl.styles import Alignment

# ========== CONFIGURATION ==========
EXCEL1_PATH = "data/excel.xlsx"
EXCEL1_SHEET = "Sheet1"
EXCEL1_COL_EL = 1
EXCEL1_COL_ID = 2
EXCEL1_COL_OTP = 3

EXCEL2_PATH = "data/serijski.xlsx"
EXCEL2_SHEET = "Sheet1"
EXCEL2_COL_OTP = 2
EXCEL2_COL_SER = 5

INPUT_DIR = "in"
OUTPUT_DIR = "out"
DEFAULT_ROW_HEIGHT = 30  # points, enough for 2-3 lines
# ===================================

def clean_el(el):
    """Remove all spaces from EL string."""
    return re.sub(r'\s+', '', el.strip())

def parse_el_otp_pair(el_cell, otp_cell):
    if pd.isna(el_cell) or pd.isna(otp_cell):
        return [], False
    el_str = str(el_cell).strip()
    otp_str = str(otp_cell).strip()
    # Split EL by newline, then clean each (remove all spaces)
    el_list = [clean_el(e) for e in re.split(r'[\n\r]+', el_str) if e.strip()]
    otp_lines = [line.strip() for line in re.split(r'[\n\r]+', otp_str) if line.strip()]
    otp_numbers = []
    for line in otp_lines:
        parts = line.split()
        if parts:
            otp_numbers.append(parts[0].strip())
        else:
            otp_numbers.append('')
    otp_list = [o for o in otp_numbers if o]
    if not el_list or not otp_list:
        return [], False
    if len(el_list) == 1 and len(otp_list) == 1:
        return [(el_list[0], otp_list[0])], False
    if len(el_list) == len(otp_list):
        return [(el_list[i], otp_list[i]) for i in range(len(el_list))], False
    return [], True

def load_second_excel():
    try:
        df = pd.read_excel(EXCEL2_PATH, sheet_name=EXCEL2_SHEET, header=None, dtype=str, engine='openpyxl')
    except Exception as e:
        print(f"Error reading second Excel: {e}")
        return {}
    otp_to_ser = defaultdict(list)
    for _, row in df.iterrows():
        otp_cell = row[EXCEL2_COL_OTP] if len(row) > EXCEL2_COL_OTP else None
        ser_cell = row[EXCEL2_COL_SER] if len(row) > EXCEL2_COL_SER else None
        if pd.isna(otp_cell) or pd.isna(ser_cell):
            continue
        otp_raw = str(otp_cell).strip()
        otp_clean = otp_raw.lstrip('0')
        ser = str(ser_cell).strip()
        otp_to_ser[otp_clean].append(ser)
    return otp_to_ser

def load_first_excel(otp_to_ser):
    try:
        df = pd.read_excel(EXCEL1_PATH, sheet_name=EXCEL1_SHEET, header=None, dtype=str, engine='openpyxl')
    except Exception as e:
        print(f"Error reading first Excel: {e}")
        return [], []
    mapping = []
    flagged_rows = []
    for idx, row in df.iterrows():
        el_cell = row[EXCEL1_COL_EL] if len(row) > EXCEL1_COL_EL else None
        id_cell = row[EXCEL1_COL_ID] if len(row) > EXCEL1_COL_ID else None
        otp_cell = row[EXCEL1_COL_OTP] if len(row) > EXCEL1_COL_OTP else None
        if pd.isna(id_cell):
            continue
        id_val = str(id_cell).strip()
        pairs, flagged = parse_el_otp_pair(el_cell, otp_cell)
        if flagged:
            otp_str = str(otp_cell).strip() if not pd.isna(otp_cell) else ''
            otp_lines = [line.strip() for line in re.split(r'[\n\r]+', otp_str) if line.strip()]
            otp_numbers = []
            for line in otp_lines:
                parts = line.split()
                if parts:
                    otp_numbers.append(parts[0].strip())
            unique_otps = list(dict.fromkeys(otp_numbers))
            ser_list = []
            for otp in unique_otps:
                ser_list.extend(otp_to_ser.get(otp, []))
            ser_list = list(dict.fromkeys(ser_list))
            flagged_rows.append((idx+1, el_cell, id_val, otp_cell, ser_list))
        else:
            for el, otp in pairs:
                mapping.append((el, id_val, otp))
    return mapping, flagged_rows

def write_dataframe_to_excel_with_formatting(df, filepath, sheet_name="Sheet1"):
    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name

    for col_idx, col_name in enumerate(df.columns, 1):
        ws.cell(row=1, column=col_idx, value=col_name)

    for row_idx, row in enumerate(df.itertuples(index=False), start=2):
        for col_idx, value in enumerate(row, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.alignment = Alignment(wrap_text=True, vertical='top')

    for col_idx, col_name in enumerate(df.columns, 1):
        col_letter = get_column_letter(col_idx)
        max_len = len(col_name)
        for row_idx in range(2, len(df) + 2):
            cell_value = ws.cell(row=row_idx, column=col_idx).value
            if cell_value:
                max_len = max(max_len, len(str(cell_value)))
        ws.column_dimensions[col_letter].width = min(max_len + 2, 50)

    for row_idx in range(1, len(df) + 2):
        ws.row_dimensions[row_idx].height = DEFAULT_ROW_HEIGHT

    wb.save(filepath)

def parse_docx_filename(filename):
    pattern = r'^G([12])\s+ELO\s+([0-9]+[a-z]?)\s*[-–]\s*([0-9]+)\.docx$'
    m = re.match(pattern, filename, re.IGNORECASE)
    if m:
        return int(m.group(1)), m.group(2), m.group(3)
    return None, None, None

def set_cell_text(cell, text, font_name='Arial', font_size=11):
    cell.text = ''
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    run = paragraph.add_run(str(text) if text is not None else '')
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def clone_row(table, source_row_idx):
    from copy import deepcopy
    source_row = table.rows[source_row_idx]
    tr = source_row._tr
    new_tr = deepcopy(tr)
    table._tbl.append(new_tr)
    return table.rows[-1]

def fill_serial_numbers(doc, serijski_list):
    if len(doc.tables) < 2:
        print("  Warning: Document has fewer than 2 tables.")
        return False
    table = doc.tables[1]
    header = table.rows[0]
    ser_col_idx = None
    for idx, cell in enumerate(header.cells):
        if "serijski broj" in cell.text.strip().lower():
            ser_col_idx = idx
            break
    if ser_col_idx is None:
        print("  Warning: Column 'Serijski broj' not found.")
        return False
    data_start_idx = 1
    if len(table.rows) <= data_start_idx:
        table.add_row()
    needed = len(serijski_list)
    current_data_rows = len(table.rows) - 1
    if needed > current_data_rows:
        for _ in range(needed - current_data_rows):
            clone_row(table, data_start_idx)
    for i, ser in enumerate(serijski_list):
        row = table.rows[data_start_idx + i]
        set_cell_text(row.cells[ser_col_idx], ser)
    return True

def main():
    Path(OUTPUT_DIR).mkdir(parents=True, exist_ok=True)

    # Step 1: Load second Excel
    print("Loading second Excel...")
    otp_to_ser = load_second_excel()
    print(f"  Loaded {len(otp_to_ser)} unique otpremnice with serials.")

    # Step 2: Load first Excel
    print("\nLoading first Excel...")
    el_id_otp, flagged_rows = load_first_excel(otp_to_ser)
    print(f"  Loaded {len(el_id_otp)} (EL, ID, otpremnica) entries.")
    if flagged_rows:
        print(f"  Flagged {len(flagged_rows)} rows (mismatched EL/otpremnica counts).")

    # Step 3: Join data for non‑flagged rows
    print("\nJoining data...")
    joined = []
    missing_otp = []   # list of (el, id, otp) that have no serials
    unique_otp_in_el_id_otp = set()
    for el, id_val, otp in el_id_otp:
        unique_otp_in_el_id_otp.add(otp)
        ser_list = otp_to_ser.get(otp, [])
        if not ser_list:
            missing_otp.append((el, id_val, otp))
        for ser in ser_list:
            joined.append((el, id_val, otp, ser))
    
    matched_otp = unique_otp_in_el_id_otp - {otp for (_, _, otp) in missing_otp}
    print(f"  Total unique otpremnice in non‑flagged rows: {len(unique_otp_in_el_id_otp)}")
    print(f"    - Matched in second Excel: {len(matched_otp)}")
    print(f"    - Missing in second Excel: {len(missing_otp)}")
    
    if missing_otp:
        print("\n  === OTpremnice WITHOUT serijski (console log) ===")
        for el, id_val, otp in missing_otp:
            print(f"    EL={el}, ID={id_val}, Otpremnica={otp}")
    else:
        print("  All otpremnice from non‑flagged rows have serials.")
    
    # Also report flagged rows summary
    if flagged_rows:
        flagged_otp_set = set()
        for _, _, _, otp_cell, _ in flagged_rows:
            if pd.notna(otp_cell):
                otp_str = str(otp_cell).strip()
                otp_lines = [line.strip() for line in re.split(r'[\n\r]+', otp_str) if line.strip()]
                for line in otp_lines:
                    parts = line.split()
                    if parts:
                        flagged_otp_set.add(parts[0].strip())
        print(f"\n  Flagged rows contain {len(flagged_otp_set)} unique otpremnice (excluded from join).")
    
    # Step 4: Save joined.xlsx with formatting
    joined_df = pd.DataFrame(joined, columns=['EL', 'ID', 'Otpremnica', 'Serijski'])
    joined_path = os.path.join(OUTPUT_DIR, "joined.xlsx")
    write_dataframe_to_excel_with_formatting(joined_df, joined_path, "Joined")
    print(f"\nSaved joined data to {joined_path}")

    # Step 5: Save flagged rows with serials (formatted)
    if flagged_rows:
        flagged_data = []
        for row_num, el_cell, id_val, otp_cell, ser_list in flagged_rows:
            ser_str = '\n'.join(ser_list) if ser_list else ''
            flagged_data.append({
                'Row': row_num,
                'EL': el_cell,
                'ID': id_val,
                'Otpremnica': otp_cell,
                'Serijski': ser_str
            })
        flagged_df = pd.DataFrame(flagged_data)
        flagged_path = os.path.join(OUTPUT_DIR, "flagged_rows.xlsx")
        write_dataframe_to_excel_with_formatting(flagged_df, flagged_path, "Flagged")
        print(f"Flagged rows saved to {flagged_path}")
        print("\n" + "="*60)
        print("FLAGGED ROWS (mismatched EL/otpremnica counts):")
        for row in flagged_data:
            print(f"  Row {row['Row']}: EL='{row['EL']}', ID='{row['ID']}', OTP='{row['Otpremnica']}'")
        print("="*60)

    # Step 6: Process DOCX files with improved messaging
    print("\nProcessing DOCX files...")
    if not os.path.isdir(INPUT_DIR):
        print(f"  Input directory '{INPUT_DIR}' not found, skipping.")
    else:
        docx_files = [f for f in os.listdir(INPUT_DIR) if f.endswith('.docx')]
        print(f"  Found {len(docx_files)} DOCX files.")
        
        # Build lookup sets (EL already cleaned in mapping)
        el_id_has_otp = set((el, id_val) for (el, id_val, _) in el_id_otp)
        el_id_has_serial = set((el, id_val) for (el, id_val, _, _) in joined)
        el_id_to_ser = defaultdict(list)
        for el, id_val, otp, ser in joined:
            el_id_to_ser[(el, id_val)].append(ser)
        
        processed = 0
        for filename in docx_files:
            group, el, id_val = parse_docx_filename(filename)
            if group is None:
                print(f"  SKIP: {filename} (invalid name format)")
                continue
            key = (el, id_val)  # EL from filename has no spaces
            if key not in el_id_has_otp:
                print(f"  SKIP: {filename} (no otpremnica found in first Excel for EL={el}, ID={id_val})")
                continue
            if key not in el_id_has_serial:
                print(f"  SKIP: {filename} (otpremnica found but no serial numbers in second Excel for EL={el}, ID={id_val})")
                continue
            ser_list = el_id_to_ser.get(key, [])
            if not ser_list:
                print(f"  SKIP: {filename} (no serial numbers found for EL={el}, ID={id_val})")
                continue
            doc_path = os.path.join(INPUT_DIR, filename)
            doc = Document(doc_path)
            if fill_serial_numbers(doc, ser_list):
                out_name = f"G{group} ELO {el}-{id_val}.docx"
                out_path = os.path.join(OUTPUT_DIR, out_name)
                doc.save(out_path)
                print(f"  OK: {out_name} ({len(ser_list)} serial(s))")
                processed += 1
            else:
                print(f"  FAIL: {filename} (table error)")
        print(f"  Processed {processed} DOCX files.")

    print("\nDone.")

if __name__ == "__main__":
    main()