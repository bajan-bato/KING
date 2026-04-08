#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import glob
import re
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from pathlib import Path

# ==================== CONFIGURATION ====================
DATA_DIR = "data"
IN_DIR = "in"
OUT_DIR = "out"
EXCEL_FILE = os.path.join(DATA_DIR, "excel.xlsx")
SERIJSKI_FILE = os.path.join(DATA_DIR, "serijski.xlsx")

EXCEL_SHEET_NAME = "Sheet1"
EXCEL_COL_EL = 1
EXCEL_COL_ID = 2
EXCEL_COL_CODE = 3

SERIJSKI_SHEET_NAME = "Sheet1"
SERIJSKI_COL_CODE = 2
SERIJSKI_COL_SERIJSKI = 5

PAD_CODE_WITH = "00"
OUTPUT_SUMMARY = "joined_data.xlsx"
# =========================================================

Path(OUT_DIR).mkdir(parents=True, exist_ok=True)

def normalize_string(s):
    return str(s).strip().lower() if pd.notna(s) else ""

def set_cell_text(cell, text, font_name='Arial', font_size=11):
    cell.text = ''
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    run = paragraph.add_run(str(text) if text is not None else '')
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def split_and_pair(el_str, code_str):
    el_list = [e.strip() for e in str(el_str).split('\n') if e.strip()]
    raw_codes = [c.strip() for c in str(code_str).split('\n') if c.strip()]
    code_list = []
    for raw in raw_codes:
        parts = raw.split()
        if parts:
            code_list.append(parts[0])
    if not code_list:
        return []
    if len(code_list) == 1:
        code_single = code_list[0]
        return [(el, code_single) for el in el_list]
    pairs = []
    for i in range(min(len(el_list), len(code_list))):
        pairs.append((el_list[i], code_list[i]))
    return pairs

def extract_el_id_from_docx(docx_path):
    try:
        doc = Document(docx_path)
        full_text = '\n'.join([para.text for para in doc.paragraphs])
        pattern = r'Evidencijska lista\s+([0-9]+[a-z]?)\s*[-–]\s*([0-9]+)'
        match = re.search(pattern, full_text, re.IGNORECASE)
        if match:
            return match.group(1).strip(), match.group(2).strip()
    except Exception:
        pass
    return None, None

# ==================== 1. READ EXCEL FILES ====================
try:
    df_excel = pd.read_excel(EXCEL_FILE, sheet_name=EXCEL_SHEET_NAME, header=None, dtype=str, engine='openpyxl')
except Exception as e:
    print(f"Error reading {EXCEL_FILE}: {e}")
    exit(1)

# Store mapping: (EL, ID) -> code, and also EL -> code for rows without ID
el_id_to_code = {}
el_to_code = {}   # first code for each EL (used when ID missing)

for idx, row in df_excel.iterrows():
    el_raw = row[EXCEL_COL_EL] if pd.notna(row[EXCEL_COL_EL]) else ""
    id_raw = row[EXCEL_COL_ID] if pd.notna(row[EXCEL_COL_ID]) else ""
    code_raw = row[EXCEL_COL_CODE] if pd.notna(row[EXCEL_COL_CODE]) else ""
    pairs = split_and_pair(el_raw, code_raw)
    for el, code in pairs:
        if not el or not code:
            continue
        norm_el = normalize_string(el)
        norm_id = normalize_string(id_raw) if id_raw else None
        el_id_to_code[(norm_el, norm_id)] = code
        if norm_el not in el_to_code:
            el_to_code[norm_el] = code

print(f"[Loaded]: {len(el_id_to_code)} (EL,ID) entries and {len(el_to_code)} unique ELs from excel.xlsx")

try:
    df_ser = pd.read_excel(SERIJSKI_FILE, sheet_name=SERIJSKI_SHEET_NAME, header=None, dtype=str, engine='openpyxl')
except Exception as e:
    print(f" [ERROR]: reading {SERIJSKI_FILE}: {e}")
    exit(1)

serijski_map = {}
for idx, row in df_ser.iterrows():
    c_val = row[SERIJSKI_COL_CODE] if pd.notna(row[SERIJSKI_COL_CODE]) else ""
    f_val = row[SERIJSKI_COL_SERIJSKI] if pd.notna(row[SERIJSKI_COL_SERIJSKI]) else ""
    if c_val and f_val:
        key = normalize_string(c_val)
        serijski_map.setdefault(key, []).append(f_val.strip())

print(f"[Loaded]: {len(serijski_map)} unique code entries from serijski.xlsx")

# ==================== 2. PROCESS ALL .DOCX FILES ====================
docx_files = glob.glob(os.path.join(IN_DIR, "*.docx"))
summary_data = []
successful_files = set()
failed = []   # (filename, reason)

for docx_path in docx_files:
    docx_name = os.path.basename(docx_path)
    
    # First try to extract EL and ID from the document content
    el, id_ = extract_el_id_from_docx(docx_path)
    if not el or not id_:
        failed.append((docx_name, "No 'Evidencijska lista' pattern found"))
        continue
    
    norm_el = normalize_string(el)
    norm_id = normalize_string(id_)
    
    # Look up code: first try exact (EL, ID), then fallback to EL only
    code = None
    if (norm_el, norm_id) in el_id_to_code:
        code = el_id_to_code[(norm_el, norm_id)]
    elif norm_el in el_to_code:
        code = el_to_code[norm_el]
        print(f"  → Using EL-only fallback for {norm_el} (ID {norm_id} not in Excel)")
    else:
        failed.append((docx_name, f"No Excel row for EL={norm_el}"))
        continue
    
    # Get serial numbers for the code
    padded_code = PAD_CODE_WITH + code
    normalized_padded = normalize_string(padded_code)
    matching_serijski = serijski_map.get(normalized_padded, [])
    if not matching_serijski:
        failed.append((docx_name, f"No serijski match for code {code}"))
        continue
    
    # Modify document
    try:
        doc = Document(docx_path)
        if len(doc.tables) < 2:
            failed.append((docx_name, "Document has fewer than 2 tables"))
            continue
        second_table = doc.tables[1]
        header_row = second_table.rows[0]
        serijski_col_idx = None
        for idx, cell in enumerate(header_row.cells):
            if normalize_string(cell.text) == "serijski broj":
                serijski_col_idx = idx
                break
        if serijski_col_idx is None:
            failed.append((docx_name, "Column 'Serijski broj' not found in second table"))
            continue
        
        num_data_rows = len(second_table.rows) - 1
        for row_idx in range(num_data_rows):
            if row_idx < len(matching_serijski):
                cell = second_table.rows[row_idx + 1].cells[serijski_col_idx]
                set_cell_text(cell, matching_serijski[row_idx])
        
        # Determine output filename: preserve the original group (G1 or G2) from input filename
        group_match = re.match(r'(G[12])', docx_name, re.IGNORECASE)
        group = group_match.group(1) if group_match else "G1"
        out_filename = f"{group} ELO {norm_el}-{norm_id}.docx"
        out_path = os.path.join(OUT_DIR, out_filename)
        doc.save(out_path)
        successful_files.add(docx_name)
        print(f"[Saved]: {out_path}")
        
        for serijski_broj in matching_serijski:
            summary_data.append({
                "A": norm_el,
                "B": norm_id,
                "C": code,
                "D": serijski_broj
            })
    except Exception as e:
        failed.append((docx_name, f"Exception: {str(e)}"))

# ==================== 3. SUMMARY ====================
print("\n" + "="*60)
print("SUMMARY")
print("="*60)
print(f"Total .docx files in '{IN_DIR}':   {len(docx_files)}")
print(f"  [SUCCESS]:        {len(successful_files)}")
print(f"  [FAILED]:         {len(failed)}")
print("="*60)

if failed:
    print("\n[FAILED FILES]:")
    for fname, reason in failed:
        print(f"  • {fname} -> {reason}")

# ==================== 4. CREATE SUMMARY XLSX ====================
if summary_data:
    df_summary = pd.DataFrame(summary_data)
    summary_path = os.path.join(OUT_DIR, OUTPUT_SUMMARY)
    df_summary.to_excel(summary_path, index=False, engine='openpyxl')
    print(f"\n  [SUMMARY]: {summary_path}")
else:
    print("\nNo data to write to summary xlsx.")