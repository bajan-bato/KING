#!/usr/bin/env python3
"""
Simple script: for every .docx in 'in/', set all cells in the "Serijski broj" column
of the second table to "n.p." (Arial 11). Output goes to 'out/' with the same name.
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

INPUT_DIR = "in"
OUTPUT_DIR = "out"
FONT_NAME = "Arial"
FONT_SIZE = 11
REPLACEMENT_TEXT = "n.p."

def set_cell_text(cell, text):
    """Replace cell content with given text, using Arial 11."""
    cell.text = ''
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE)
    # Ensure East Asian characters also use the same font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

def process_docx(input_path, output_path):
    """Open docx, replace serial numbers with n.p., save to output path."""
    doc = Document(input_path)

    # Need at least two tables
    if len(doc.tables) < 2:
        print(f"  SKIP: {os.path.basename(input_path)} (less than 2 tables)")
        return False

    table = doc.tables[1]          # second table (index 1)
    if len(table.rows) == 0:
        print(f"  SKIP: {os.path.basename(input_path)} (second table has no rows)")
        return False

    # Find column index of "Serijski broj" in the header row
    header_row = table.rows[0]
    ser_col_idx = None
    for idx, cell in enumerate(header_row.cells):
        if "serijski broj" in cell.text.strip().lower():
            ser_col_idx = idx
            break

    if ser_col_idx is None:
        print(f"  SKIP: {os.path.basename(input_path)} (column 'Serijski broj' not found)")
        return False

    # Replace every data row (skip header)
    for row_idx in range(1, len(table.rows)):
        cell = table.rows[row_idx].cells[ser_col_idx]
        set_cell_text(cell, REPLACEMENT_TEXT)

    doc.save(output_path)
    return True

def main():
    # Create output directory if it doesn't exist
    Path(OUTPUT_DIR).mkdir(parents=True, exist_ok=True)

    if not os.path.isdir(INPUT_DIR):
        print(f"Input directory '{INPUT_DIR}' not found. Exiting.")
        return

    docx_files = [f for f in os.listdir(INPUT_DIR) if f.lower().endswith('.docx')]
    print(f"Found {len(docx_files)} DOCX file(s).")

    processed = 0
    for filename in docx_files:
        in_path = os.path.join(INPUT_DIR, filename)
        out_path = os.path.join(OUTPUT_DIR, filename)
        print(f"Processing: {filename} ... ", end='')
        if process_docx(in_path, out_path):
            print("OK")
            processed += 1
        else:
            print("SKIPPED")

    print(f"\nDone. Processed {processed} file(s).")

if __name__ == "__main__":
    main()